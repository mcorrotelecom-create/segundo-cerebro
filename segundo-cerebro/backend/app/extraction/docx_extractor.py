"""Extracción de texto de archivos Word: párrafos y tablas, en orden."""
from app.extraction import ExtractedUnit


def extract_docx(file_path: str) -> list[ExtractedUnit]:
    import docx

    document = docx.Document(file_path)
    units: list[ExtractedUnit] = []

    for i, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if text:
            units.append(ExtractedUnit(text=text, source_ref={"paragraph": i}))

    for t_index, table in enumerate(document.tables, start=1):
        rows_txt = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows_txt.append(" | ".join(cells))
        if rows_txt:
            units.append(ExtractedUnit(text="\n".join(rows_txt), source_ref={"table": t_index}))

    if not units:
        units.append(ExtractedUnit(text="", source_ref={"note": "documento_sin_texto_extraible"}))
    return units
